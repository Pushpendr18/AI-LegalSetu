#------------------------------------------------------------------------------------
import os
import logging
import json
from io import BytesIO
from typing import List, Optional
from datetime import datetime
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
import uvicorn
from auth import router as auth_router
import sqlite3
import uuid

import google.generativeai as genai
from pdfminer.high_level import extract_text
import docx
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss
from rag_service import legal_rag
import PyPDF2
import io

# CORS setup moved below after FastAPI app initialization to avoid referencing `app` before it's defined.

# SQLite Database setup
def get_db_connection():
    conn = sqlite3.connect('legal_chatbot.db')
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS chat_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT UNIQUE,
            title TEXT,
            messages TEXT,
            created_at TEXT,
            updated_at TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_db()


CONTRACT_TEMPLATES = {
    "nda": {
        "name": "Non-Disclosure Agreement",
        "template": """
THIS NON-DISCLOSURE AGREEMENT (the "Agreement") is made on {date}

BETWEEN:
{party1_name}, having address at {party1_address} ("Disclosing Party")

AND:
{party2_name}, having address at {party2_address} ("Receiving Party")

1. CONFIDENTIAL INFORMATION
The term "Confidential Information" shall mean {confidential_info}.

2. OBLIGATIONS
The Receiving Party shall:
(a) Maintain the confidentiality of the Information
(b) Use the same degree of care to protect Confidential Information  
(c) Not disclose Confidential Information to any third party

3. TERM
This Agreement shall remain in effect for {duration} years.

IN WITNESS WHEREOF, the parties have executed this Agreement.

_________________________
{party1_name}

_________________________
{party2_name}
        """,
        "fields": ["party1_name", "party1_address", "party2_name", "party2_address", 
                  "confidential_info", "duration"]
    },
    
    "rental_agreement": {
        "name": "Residential Rental Agreement", 
        "template": """
RESIDENTIAL RENTAL AGREEMENT

This Agreement made on {date} between:

LANDLORD: {landlord_name}, address: {landlord_address}

TENANT: {tenant_name}, address: {tenant_address}

PROPERTY: {property_address}

TERMS:
1. Rent: ₹{rent_amount} per month, payable on {rent_due_date}
2. Security Deposit: ₹{security_deposit}
3. Lease Term: {lease_months} months from {start_date} to {end_date}
4. Utilities: {utilities_responsibility}

RULES:
- No structural changes without permission
- Proper maintenance required
- {pet_policy}

IN WITNESS WHEREOF, the parties execute this Agreement.

_________________________
{landlord_name}

_________________________
{tenant_name}
        """,
        "fields": ["landlord_name", "landlord_address", "tenant_name", "tenant_address",
                  "property_address", "rent_amount", "security_deposit", "lease_months",
                  "start_date", "end_date", "utilities_responsibility", "pet_policy"]
    },
    
    "service_agreement": {
        "name": "Service Agreement",
        "template": """
SERVICE AGREEMENT

This Service Agreement is made on {date} between:

SERVICE PROVIDER: {provider_name}, address: {provider_address}

CLIENT: {client_name}, address: {client_address}

SCOPE OF SERVICES:
{services_description}

TERMS:
1. Service Fee: ₹{service_fee}
2. Payment Terms: {payment_terms}
3. Term: {contract_term}
4. Termination: {termination_clause}

WARRANTY:
{service_warranty}

IN WITNESS WHEREOF, the parties execute this Agreement.

_________________________
{provider_name}

_________________________
{client_name}
        """,
        "fields": ["provider_name", "provider_address", "client_name", "client_address",
                  "services_description", "service_fee", "payment_terms", "contract_term",
                  "termination_clause", "service_warranty"]
    }
}
# Pydantic Models
class ChatMessage(BaseModel):
    sender: str
    text: str
    timestamp: str

class ChatSession(BaseModel):
    session_id: str
    title: str
    messages: List[ChatMessage]
    created_at: str
    updated_at: str

# ---------- Load .env ----------
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "models/gemini-2.5-flash")

if not GEMINI_API_KEY:
    raise RuntimeError("❌ GEMINI_API_KEY not found — add it to your .env file")

genai.configure(api_key=GEMINI_API_KEY)

# ------------------------------------- Logging ----------------------------------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("ai-legal-assistant")

# ---------- FastAPI ----------
app = FastAPI(title="AI Legal Assistant - Gemini Version")
app.include_router(auth_router)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- Utility functions ----------
def extract_text_from_pdf_bytes(content: bytes) -> str:
    """Extracts text safely from PDF file bytes."""
    try:
        text = extract_text(BytesIO(content))
        if not text.strip():
            raise ValueError("No text found in PDF.")
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting PDF text: {e}")

def build_faiss_index(text: str):
    """Break text into chunks and build FAISS index for semantic search."""
    embed_model = SentenceTransformer("all-MiniLM-L6-v2")
    chunks = [text[i:i+1000] for i in range(0, len(text), 1000)]
    embeddings = embed_model.encode(chunks, convert_to_numpy=True)
    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings.astype("float32"))
    return embed_model, index, chunks

# ---------- Gemini Helper Function ----------
def call_gemini_direct(prompt: str, model=GEMINI_MODEL):
    """Direct call to Gemini without chat memory"""
    try:
        model_instance = genai.GenerativeModel(model)
        response = model_instance.generate_content(prompt)
        return response.text
    except Exception as e:
        logger.error(f"Gemini direct call error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ---------- Pydantic Models ----------
class AskQueryRequest(BaseModel):
    query: str

class AskQueryResponse(BaseModel):
    answer: str
    sources: Optional[List[str]] = None

# ---------- Gemini Chat with Memory Support ----------
chat_sessions = {}

def call_gemini_chat_with_memory(user_id: str, user_message: str, model=GEMINI_MODEL):
    """Maintain chat context for a user using Gemini chat sessions."""
    try:
        if user_id not in chat_sessions:
            model_instance = genai.GenerativeModel(model)
            chat_sessions[user_id] = model_instance.start_chat(history=[])
        chat = chat_sessions[user_id]
        response = chat.send_message(user_message)
        return response.text
    except Exception as e:
        logger.error(f"Gemini chat memory error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ---------- Routes ----------
@app.get("/")
def home():
    return {"message": "✅ Gemini Legal Assistant Backend is running"}

# STEP 2: main.py mein routes ke section mein yeh add karo
@app.post("/api/generate-contract")
async def generate_contract(contract_data: dict):
    try:
        template_type = contract_data.get("template_type")
        form_data = contract_data.get("form_data", {})
        
        if template_type not in CONTRACT_TEMPLATES:
            raise HTTPException(status_code=400, detail="Invalid template type")
        
        template = CONTRACT_TEMPLATES[template_type]["template"]
        
        # Add current date if not provided
        if "date" not in form_data:
            form_data["date"] = datetime.now().strftime("%B %d, %Y")
            
        # Add default values for optional fields
        if template_type == "rental_agreement":
            form_data.setdefault("rent_due_date", "the 1st of each month")
            form_data.setdefault("end_date", "as per lease term")
            form_data.setdefault("utilities_responsibility", "Tenant shall pay all utilities")
            form_data.setdefault("pet_policy", "No pets allowed without written permission")
            
        elif template_type == "service_agreement":
            form_data.setdefault("payment_terms", "50% advance, 50% on completion")
            form_data.setdefault("termination_clause", "Either party may terminate with 30 days written notice")
            form_data.setdefault("service_warranty", "Services will be performed in a professional manner")
        
        # Replace placeholders with actual data
        generated_contract = template.format(**form_data)
        
        return {
            "success": True,
            "contract": generated_contract,
            "template_name": CONTRACT_TEMPLATES[template_type]["name"]
        }
        
    except KeyError as e:
        raise HTTPException(status_code=400, detail=f"Missing required field: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating contract: {str(e)}")

@app.post("/api/ask-query", response_model=AskQueryResponse)
async def ask_query(req: AskQueryRequest):
    """Maintains memory across chats per user (context-aware)."""
    try:
        user_id = "default_user"
        logger.info(f"ask-query received from {user_id}")
        answer = call_gemini_chat_with_memory(user_id, req.query)
        return {"answer": answer, "sources": []}
    except Exception as e:
        logger.error(f"Error in /api/ask-query: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/summarize")
async def summarize(file: UploadFile = File(...)):
    """Uploads a PDF/DOCX/TXT and returns a structured summary."""
    try:
        content = await file.read()
        filename = file.filename.lower()

        if filename.endswith(".pdf"):
            text = extract_text_from_pdf_bytes(content)
        elif filename.endswith(".docx"):
            doc = docx.Document(BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs])
        elif filename.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(400, "Unsupported file type. Please use PDF, DOCX, or TXT.")

        chunks = [text[i:i+6000] for i in range(0, len(text), 6000)]
        summaries = []
        model = genai.GenerativeModel(GEMINI_MODEL)
        for i, chunk in enumerate(chunks):
            prompt = f"Summarize the following legal document part {i+1}/{len(chunks)}:\n\n{chunk}"
            response = model.generate_content(prompt)
            summaries.append(response.text)

        final_prompt = (
            "Combine these partial summaries into a final structured legal summary "
            "with sections like: 'Overview', 'Key Clauses', 'Risks', and 'Recommendations'.\n\n"
            + "\n\n---\n\n".join(summaries)
        )
        final_summary = model.generate_content(final_prompt).text
        return {"summary": final_summary}
    except Exception as e:
        logger.error(f"Error in /api/summarize: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/classify")
async def classify_text(file: UploadFile = File(...)):
    """Classify the uploaded legal document as Agreement, Notice, Petition, Judgment, or Other."""
    try:
        content = await file.read()
        filename = file.filename.lower()
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf_bytes(content)
        elif filename.endswith(".docx"):
            doc = docx.Document(BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs])
        elif filename.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(400, "Unsupported file type.")

        prompt = f"""
Classify the following document into one of:
- Agreement
- Petition
- Notice
- Judgment
- Other

Respond only with pure JSON like: {{"category": "...", "confidence": 0.XX"}}

Document Text:
{text[:6000]}
"""
        model = genai.GenerativeModel(GEMINI_MODEL)
        response = model.generate_content(prompt)
        import re
        clean = re.sub(r"```json|```", "", response.text.strip()).strip()
        try:
            parsed = json.loads(clean)
        except Exception:
            parsed = {"category": "Other", "confidence": 0.0}
        return parsed
    except Exception as e:
        logger.error(f"Error in /api/classify: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/ask-doc-query")
async def ask_doc_query(file: UploadFile = File(...), query: str = Form(...)):
    """Ask a question based on document context using RAG (semantic search + Gemini)."""
    try:
        content = await file.read()
        filename = file.filename.lower()
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf_bytes(content)
        elif filename.endswith(".docx"):
            doc = docx.Document(BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs])
        elif filename.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(400, "Unsupported file type (use pdf/docx/txt).")

        embed_model, index, chunks = build_faiss_index(text)
        q_emb = embed_model.encode([query], convert_to_numpy=True).astype("float32")
        D, I = index.search(q_emb, 3)
        top_context = "\n\n".join([chunks[i] for i in I[0]])

        prompt = f"""
You are LegalSetu, an AI legal assistant. 
Use the document context below to answer the user's question accurately and concisely.
If information is not found in context, say "Not found in document."

Context:
{top_context}

Question: {query}
"""
        g_model = genai.GenerativeModel(GEMINI_MODEL)
        response = g_model.generate_content(prompt)
        return {"answer": response.text}
    except Exception as e:
        logger.error(f"Error in /api/ask-doc-query (RAG): {e}")
        raise HTTPException(status_code=500, detail=str(e))

# 🚨 RAG ENDPOINTS - YAHAN SE ADD KARO 🚨

@app.post("/api/add-to-knowledge")
async def add_to_knowledge(file: UploadFile = File(...)):
    """Add uploaded file to RAG knowledge base"""
    try:
        content = await file.read()
        filename = file.filename.lower()
        
        # Extract text based on file type
        if filename.endswith('.pdf'):
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        elif filename.endswith('.docx'):
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            # For txt files
            text = content.decode('utf-8', errors='ignore')
        
        # Add to RAG system
        doc_count = legal_rag.add_document(
            text, 
            metadata={"source": file.filename, "type": "legal_document"}
        )
        
        return {
            "message": f"Successfully added {doc_count} document chunks to knowledge base",
            "chunks_added": doc_count
        }
        
    except Exception as e:
        logger.error(f"Error in /api/add-to-knowledge: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.post("/api/rag-chat")
async def rag_chat(request: dict):
    """Enhanced chat with RAG context"""
    try:
        query = request.get("query", )
        user_id = "default_user"  # Same as your existing chat
        
        # Get relevant context from knowledge base
        context = legal_rag.get_context_for_query(query)
        
        # Enhanced prompt with context
        enhanced_prompt = f"""
        You are LegalSetu, an AI legal assistant. 
        Based on the following legal context from uploaded documents, please answer the user's question accurately and helpfully.

        RELEVANT LEGAL CONTEXT:
        {context}

        USER QUESTION: {query}

        Please provide a comprehensive legal answer based on the context above. 
        If the context doesn't contain relevant information, use your general legal knowledge but indicate this clearly.
        Always cite relevant sources when possible.
        """
        
        # Use your existing Gemini chat with memory
        answer = call_gemini_chat_with_memory(user_id, enhanced_prompt)
        
        # Also get the source documents for citations
        similar_docs = legal_rag.search_similar(query)
        sources = [{"content": doc.page_content[:200] + "...", "source": doc.metadata.get('source', 'Unknown')} 
                  for doc in similar_docs]
        
        return {
            "answer": answer,
            "sources": sources,
            "has_context": len(context) > 0
        }
        
    except Exception as e:
        logger.error(f"Error in /api/rag-chat: {e}")
        raise HTTPException(status_code=500, detail=f"RAG chat error: {str(e)}")

# ---------- Run ----------
if __name__ == "__main__":
    print("🚀 Starting Gemini Backend on http://127.0.0.1:8000 ...")
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)